import csv
import io
import os
from typing import Dict, Any, List

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from core.models import MeetingAnalysisResponse


def generate_markdown_report(report_data: MeetingAnalysisResponse) -> str:
    """Generate Markdown report from existing MeetingAnalysisResponse data."""
    md = []
    md.append(f"# {report_data.title}\n")

    if report_data.metadata:
        meta = report_data.metadata
        md.append("## Meeting Metadata")
        if meta.generation_timestamp:
            md.append(f"- **Generated At:** {meta.generation_timestamp}")
        if meta.detected_language:
            md.append(f"- **Detected Language:** {meta.detected_language}")
        if meta.transcription_engine:
            md.append(f"- **Engine Used:** {meta.transcription_engine}")
        if meta.transcript_chunks_count is not None:
            md.append(f"- **Transcript Chunks:** {meta.transcript_chunks_count}")
        if meta.duration_seconds is not None:
            md.append(f"- **Duration:** {meta.duration_seconds:.1f}s")
        if meta.processing_time_seconds is not None:
            md.append(f"- **Processing Time:** {meta.processing_time_seconds:.1f}s")
        md.append("")

    md.append("## Executive Summary")
    md.append(report_data.summary)
    md.append("")

    md.append("## Key Decisions")
    md.append(report_data.key_decisions)
    md.append("")

    md.append("## Action Items")
    md.append(report_data.action_items)
    md.append("")

    md.append("## Open Questions")
    md.append(report_data.questions)
    md.append("")

    if report_data.risks:
        md.append("## Risks")
        md.append(report_data.risks)
        md.append("")

    if report_data.next_steps:
        md.append("## Next Steps")
        md.append(report_data.next_steps)
        md.append("")

    return "\n".join(md)


def generate_action_items_csv(action_items_text: str) -> str:
    """Parse action items text and generate CSV content with Task, Owner, Deadline, Status."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Task", "Owner", "Deadline", "Status"])

    lines = [line.strip() for line in action_items_text.split("\n") if line.strip()]
    current_task = ""
    owner = "Not specified"
    deadline = "Not specified"

    for line in lines:
        if line.lower().startswith(("- task:", "task:", "-")):
            if current_task:
                writer.writerow([current_task, owner, deadline, "Pending"])
            current_task = line.lstrip("- ").replace("Task:", "").strip()
            owner = "Not specified"
            deadline = "Not specified"
        elif "owner:" in line.lower():
            owner = line.split(":", 1)[-1].strip()
        elif "deadline:" in line.lower():
            deadline = line.split(":", 1)[-1].strip()
        else:
            if not current_task:
                current_task = line
            else:
                current_task += " " + line

    if current_task and current_task.lower() != "no action items found.":
        writer.writerow([current_task, owner, deadline, "Pending"])

    return output.getvalue()


def generate_pdf_report(report_data: MeetingAnalysisResponse, output_path: str) -> str:
    """Generate PDF report file using reportlab from existing report data."""
    doc = SimpleDocTemplate(output_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#1E293B'),
        spaceAfter=12
    )

    h2_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor('#2563EB'),
        spaceBefore=10,
        spaceAfter=6
    )

    body_style = styles['BodyText']

    story.append(Paragraph(report_data.title, title_style))
    story.append(Spacer(1, 10))

    if report_data.metadata:
        story.append(Paragraph("Meeting Statistics", h2_style))
        meta = report_data.metadata
        meta_lines = []
        if meta.generation_timestamp:
            meta_lines.append(f"Generated: {meta.generation_timestamp}")
        if meta.detected_language:
            meta_lines.append(f"Language: {meta.detected_language}")
        if meta.transcription_engine:
            meta_lines.append(f"Engine: {meta.transcription_engine}")
        if meta.transcript_chunks_count:
            meta_lines.append(f"Chunks: {meta.transcript_chunks_count}")
        story.append(Paragraph("<br/>".join(meta_lines), body_style))
        story.append(Spacer(1, 10))

    sections = [
        ("Executive Summary", report_data.summary),
        ("Key Decisions", report_data.key_decisions),
        ("Action Items", report_data.action_items),
        ("Open Questions", report_data.questions),
    ]

    if report_data.risks:
        sections.append(("Risks", report_data.risks))
    if report_data.next_steps:
        sections.append(("Next Steps", report_data.next_steps))

    for title, content in sections:
        story.append(Paragraph(title, h2_style))
        formatted_content = content.replace("\n", "<br/>")
        story.append(Paragraph(formatted_content, body_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    return output_path


def generate_docx_report(report_data: MeetingAnalysisResponse, output_path: str) -> str:
    """Generate Word (.docx) report file or fallback text file if python-docx is not installed."""
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(report_data.title, 0)

        if report_data.metadata:
            doc.add_heading("Meeting Statistics", level=1)
            meta = report_data.metadata
            p = doc.add_paragraph()
            if meta.generation_timestamp:
                p.add_run(f"Generated: {meta.generation_timestamp}\n")
            if meta.detected_language:
                p.add_run(f"Language: {meta.detected_language}\n")
            if meta.transcription_engine:
                p.add_run(f"Engine: {meta.transcription_engine}\n")

        sections = [
            ("Executive Summary", report_data.summary),
            ("Key Decisions", report_data.key_decisions),
            ("Action Items", report_data.action_items),
            ("Open Questions", report_data.questions),
        ]
        if report_data.risks:
            sections.append(("Risks", report_data.risks))
        if report_data.next_steps:
            sections.append(("Next Steps", report_data.next_steps))

        for title, content in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

        doc.save(output_path)
        return output_path
    except ImportError:
        # Fallback to plain text document if python-docx is unavailable
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(generate_markdown_report(report_data))
        return output_path
