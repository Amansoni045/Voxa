from backend.models.meeting import MeetingAnalysis
from backend.exports.markdown import generate_markdown_report

def generate_docx_report(report_data: MeetingAnalysis, output_path: str) -> str:
    try:
        import docx
        doc = docx.Document()
        doc.add_heading(report_data.title, 0)

        if report_data.metadata:
            doc.add_heading("Meeting Statistics", level=1)
            meta = report_data.metadata
            p = doc.add_paragraph()
            if meta.generation_timestamp:
                p.add_run(f"Generated: {meta.generation_timestamp}\n")
            if meta.detected_language:
                p.add_run(f"Language: {meta.detected_language}\n")
            if meta.transcription_engine:
                p.add_run(f"Engine: {meta.transcription_engine}\n")

        sections = [
            ("Executive Summary", report_data.summary),
            ("Key Decisions", report_data.key_decisions),
            ("Action Items", report_data.action_items),
            ("Open Questions", report_data.questions),
        ]
        if report_data.risks:
            sections.append(("Risks", report_data.risks))
        if report_data.next_steps:
            sections.append(("Next Steps", report_data.next_steps))

        for title, content in sections:
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

        doc.save(output_path)
        return output_path
    except ImportError:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(generate_markdown_report(report_data))
        return output_path
